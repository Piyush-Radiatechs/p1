"""Extract job-description text from PDF, Word, and plain-text files."""

from io import BytesIO
from pathlib import Path

from docx import Document

from app.exceptions import DocumentExtractionError
from app.services.pdf_parser import extract_pdf_text

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from a .docx Word document, including tables."""
    if not docx_bytes:
        raise DocumentExtractionError("Uploaded file is empty.")

    try:
        document = Document(BytesIO(docx_bytes))
    except Exception as exc:
        raise DocumentExtractionError(f"Invalid or corrupted Word document: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise DocumentExtractionError("No extractable text found in the Word document.")

    return "\n".join(parts)


def extract_plain_text(file_bytes: bytes) -> str:
    """Decode a .txt or .md file as UTF-8 (with fallbacks)."""
    if not file_bytes:
        raise DocumentExtractionError("Uploaded file is empty.")

    text = None
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        raise DocumentExtractionError("Could not decode the text file.")

    stripped = text.strip()
    if not stripped:
        raise DocumentExtractionError("The text file is empty.")

    return stripped


def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch extraction based on the uploaded filename extension."""
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return extract_docx_text(file_bytes)
    if suffix in {".txt", ".md"}:
        return extract_plain_text(file_bytes)

    raise DocumentExtractionError(
        "Unsupported file type. Upload a PDF, Word (.docx), or text (.txt) file."
    )
